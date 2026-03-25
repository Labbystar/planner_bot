from __future__ import annotations

import calendar
import csv
import io
from datetime import datetime, timedelta, timezone
from io import BytesIO

from aiogram import F, Router
from aiogram.fsm.context import FSMContext
from aiogram.fsm.state import State, StatesGroup
from aiogram.types import BufferedInputFile, CallbackQuery, Message
from docx import Document
from openpyxl import Workbook, load_workbook

from app.keyboards.reminders import active_filters_kb, calendar_kb, pager, reminder_actions, settings_kb
from app.services.reminders import (
    create_reminder,
    list_active_reminders,
    list_all_reminders,
    list_assigned_to_me,
    list_by_time_window,
    list_created_by_me,
    list_month_counts,
    list_overdue,
    search_reminders,
    stats,
    count_attachments,
)
from app.services.users import get_user
from app.utils.formatting import calendar_title, list_line, page_header, reminder_card, stats_text, user_label
from app.utils.time import parse_user_time, to_local

router = Router()


class SearchFlow(StatesGroup):
    query = State()


class CsvImportFlow(StatesGroup):
    wait_file = State()


class ExcelImportFlow(StatesGroup):
    wait_file = State()


@router.message(F.text == "📋 Активные")
async def active(message: Message) -> None:
    await message.answer("Фильтр списка:", reply_markup=active_filters_kb())
    await send_reminders_page(message, message.from_user.id, 1)


@router.message(F.text == "📥 Мне поставили")
async def assigned_to_me(message: Message) -> None:
    await send_assigned_page(message, message.from_user.id, 1)


@router.message(F.text == "📤 Я поставил")
async def created_by_me(message: Message) -> None:
    await send_owner_page(message, message.from_user.id, 1)


@router.message(F.text == "📅 Сегодня")
async def today(message: Message) -> None:
    await send_window(message, "📅 Сегодня", 0, 1)


@router.message(F.text == "🌤 Завтра")
async def tomorrow(message: Message) -> None:
    await send_window(message, "🌤 Завтра", 1, 2)


@router.message(F.text == "🗓 Неделя")
async def week(message: Message) -> None:
    await send_window(message, "🗓 Неделя", 0, 7)


@router.message(F.text == "🗓 Календарь")
async def calendar_view(message: Message) -> None:
    now = datetime.now()
    await send_month_calendar(message, message.from_user.id, now.year, now.month)


@router.message(F.text == "⏳ Просроченные")
async def overdue(message: Message) -> None:
    user = await get_user(message.from_user.id)
    if not user:
        return
    reminders = await list_overdue(message.from_user.id)
    if not reminders:
        await message.answer("<b>⏳ Просроченные</b>\nПока пусто.", parse_mode="HTML")
        return
    lines = ["<b>⏳ Просроченные</b>"]
    for r in reminders[:15]:
        dt_local = to_local(datetime.fromisoformat(r["scheduled_at_utc"]), user["timezone_name"])
        owner = await get_user(r.get("owner_user_id"))
        assignee = await get_user(r.get("assigned_user_id"))
        mode = 'assigned' if r.get('assigned_user_id') == message.from_user.id and r.get('owner_user_id') != message.from_user.id else 'owner' if r.get('owner_user_id') == message.from_user.id and r.get('assigned_user_id') != message.from_user.id else 'shared'
        lines.append(list_line(r, dt_local, user_label(owner, r.get('owner_user_id')), user_label(assignee, r.get('assigned_user_id')), mode))
    await message.answer("\n\n".join(lines), parse_mode="HTML")


@router.message(F.text == "🔎 Поиск")
async def search_start(message: Message, state: FSMContext) -> None:
    await state.set_state(SearchFlow.query)
    await message.answer("Введи слово или часть текста задачи для поиска.")


@router.message(SearchFlow.query)
async def search_run(message: Message, state: FSMContext) -> None:
    user = await get_user(message.from_user.id)
    if not user:
        return
    rows = await search_reminders(message.from_user.id, message.text)
    await state.clear()
    if not rows:
        await message.answer("Ничего не нашел.")
        return
    lines = [f"<b>🔎 Поиск: {message.text}</b>"]
    for r in rows:
        dt_local = to_local(datetime.fromisoformat(r["scheduled_at_utc"]), user["timezone_name"])
        owner = await get_user(r.get("owner_user_id"))
        assignee = await get_user(r.get("assigned_user_id"))
        mode = 'assigned' if r.get('assigned_user_id') == message.from_user.id and r.get('owner_user_id') != message.from_user.id else 'owner' if r.get('owner_user_id') == message.from_user.id and r.get('assigned_user_id') != message.from_user.id else 'shared'
        lines.append(list_line(r, dt_local, user_label(owner, r.get('owner_user_id')), user_label(assignee, r.get('assigned_user_id')), mode))
    await message.answer("\n\n".join(lines), parse_mode="HTML")


@router.message(F.text == "📤 Экспорт CSV")
async def export_csv(message: Message) -> None:
    user = await get_user(message.from_user.id)
    if not user:
        return
    rows = await list_all_reminders(message.from_user.id)
    if not rows:
        await message.answer("Пока нечего экспортировать.")
        return
    output = io.StringIO()
    writer = csv.writer(output)
    writer.writerow(["id","text","note","category","priority","status","owner_user_id","assigned_user_id","scheduled_local","scheduled_utc","created_at_utc","updated_at_utc","completed_at_utc"])
    for r in rows:
        dt_local = to_local(datetime.fromisoformat(r["scheduled_at_utc"]), user["timezone_name"])
        writer.writerow([r.get("id", ""),r.get("text", ""),r.get("note", "") or "",r.get("category", "") or "",r.get("priority", "") or "",r.get("status", "") or "",r.get("owner_user_id", ""),r.get("assigned_user_id", ""),dt_local.strftime('%d.%m.%Y %H:%M'),r.get("scheduled_at_utc", "") or "",r.get("created_at", "") or "",r.get("updated_at", "") or "",r.get("completed_at", "") or ""])
    payload = output.getvalue().encode("utf-8-sig")
    filename = f"napomnime_export_{datetime.now().strftime('%Y%m%d_%H%M')}.csv"
    await message.answer_document(BufferedInputFile(payload, filename=filename), caption="Готово. Это полный экспорт задач и статусов в CSV.")


@router.message(F.text == "📤 Экспорт Excel")
async def export_excel(message: Message) -> None:
    user = await get_user(message.from_user.id)
    if not user:
        return
    rows = await list_all_reminders(message.from_user.id)
    if not rows:
        await message.answer("Пока нечего экспортировать.")
        return
    wb = Workbook()
    ws = wb.active
    ws.title = 'Tasks'
    headers = ["id","text","note","category","priority","status","owner","assignee","scheduled_local","scheduled_utc","created_at_utc","updated_at_utc","completed_at_utc"]
    ws.append(headers)
    for r in rows:
        owner = await get_user(r.get('owner_user_id'))
        assignee = await get_user(r.get('assigned_user_id'))
        dt_local = to_local(datetime.fromisoformat(r['scheduled_at_utc']), user['timezone_name'])
        ws.append([r.get('id'), r.get('text'), r.get('note') or '', r.get('category'), r.get('priority'), r.get('status'), user_label(owner, r.get('owner_user_id')), user_label(assignee, r.get('assigned_user_id')), dt_local.strftime('%Y-%m-%d %H:%M'), r.get('scheduled_at_utc'), r.get('created_at'), r.get('updated_at'), r.get('completed_at') or ''])
    bio = BytesIO()
    wb.save(bio)
    bio.seek(0)
    await message.answer_document(BufferedInputFile(bio.read(), filename=f"napomnime_export_{datetime.now().strftime('%Y%m%d_%H%M')}.xlsx"), caption="Готово. Экспорт задач в Excel.")


@router.message(F.text == "📄 Word-отчёт")
async def export_word(message: Message) -> None:
    user = await get_user(message.from_user.id)
    if not user:
        return
    rows = await list_all_reminders(message.from_user.id)
    if not rows:
        await message.answer("Пока нечего экспортировать.")
        return
    doc = Document()
    doc.add_heading('Отчёт по задачам', 1)
    doc.add_paragraph(f"Сформировано: {datetime.now().strftime('%d.%m.%Y %H:%M')}")
    for r in rows:
        owner = await get_user(r.get('owner_user_id'))
        assignee = await get_user(r.get('assigned_user_id'))
        dt_local = to_local(datetime.fromisoformat(r['scheduled_at_utc']), user['timezone_name'])
        p = doc.add_paragraph()
        p.add_run(r['text']).bold = True
        doc.add_paragraph(f"Срок: {dt_local.strftime('%d.%m.%Y %H:%M')} | Статус: {r['status']} | Исполнитель: {user_label(assignee, r.get('assigned_user_id'))} | Постановщик: {user_label(owner, r.get('owner_user_id'))}")
        if r.get('note'):
            doc.add_paragraph(f"Заметка: {r['note']}")
    bio = BytesIO()
    doc.save(bio)
    bio.seek(0)
    await message.answer_document(BufferedInputFile(bio.read(), filename=f"napomnime_report_{datetime.now().strftime('%Y%m%d_%H%M')}.docx"), caption="Готово. Word-отчёт по задачам.")


@router.message(F.text == "📥 Импорт CSV")
async def import_csv_start(message: Message, state: FSMContext) -> None:
    await state.set_state(CsvImportFlow.wait_file)
    await message.answer("Пришли CSV-файл с задачами. Обязательные колонки: text, scheduled_local. Необязательные: note, category, priority, assigned_user_id")


@router.message(CsvImportFlow.wait_file, F.document)
async def import_csv_file(message: Message, state: FSMContext) -> None:
    user = await get_user(message.from_user.id)
    if not user:
        await state.clear(); return
    if not message.document.file_name.lower().endswith('.csv'):
        await message.answer("Нужен именно CSV-файл."); return
    file = await message.bot.get_file(message.document.file_id)
    content = await message.bot.download_file(file.file_path)
    raw = content.read()
    text = raw.decode('utf-8-sig', errors='replace')
    reader = csv.DictReader(io.StringIO(text))
    imported, errors = 0, []
    for idx, row in enumerate(reader, start=2):
        try:
            reminder_text = (row.get('text') or '').strip(); when_raw = (row.get('scheduled_local') or '').strip()
            if not reminder_text or not when_raw:
                raise ValueError('нужны text и scheduled_local')
            dt_local = parse_user_time(when_raw, user['timezone_name'])
            assigned_raw = (row.get('assigned_user_id') or '').strip()
            assigned_user_id = message.from_user.id
            if assigned_raw:
                if not assigned_raw.isdigit() or not await get_user(int(assigned_raw)):
                    raise ValueError('assigned_user_id не найден среди активных пользователей')
                assigned_user_id = int(assigned_raw)
            await create_reminder(message.from_user.id, assigned_user_id, reminder_text, dt_local.astimezone(timezone.utc).isoformat(), (row.get('category') or 'work').strip() or 'work', (row.get('priority') or 'medium').strip() or 'medium', (row.get('note') or '').strip() or None)
            imported += 1
        except Exception as e:
            errors.append(f"Строка {idx}: {e}")
    await state.clear()
    await message.answer(f"Импорт завершён.\n\n✅ Импортировано: {imported}" + (f"\n⚠️ Ошибок: {len(errors)}\n" + "\n".join(errors[:10]) if errors else ""))


@router.message(F.text == "📥 Импорт Excel")
async def import_excel_start(message: Message, state: FSMContext) -> None:
    await state.set_state(ExcelImportFlow.wait_file)
    await message.answer("Пришли Excel-файл .xlsx. На первом листе должны быть колонки: text, scheduled_local, note, category, priority, assigned_user_id")


@router.message(ExcelImportFlow.wait_file, F.document)
async def import_excel_file(message: Message, state: FSMContext) -> None:
    user = await get_user(message.from_user.id)
    if not user:
        await state.clear(); return
    if not message.document.file_name.lower().endswith('.xlsx'):
        await message.answer("Нужен именно Excel-файл .xlsx."); return
    file = await message.bot.get_file(message.document.file_id)
    content = await message.bot.download_file(file.file_path)
    wb = load_workbook(BytesIO(content.read()))
    ws = wb.active
    rows = list(ws.iter_rows(values_only=True))
    if not rows:
        await message.answer('Пустой файл'); await state.clear(); return
    headers = [str(x).strip() if x is not None else '' for x in rows[0]]
    imported, errors = 0, []
    for idx, row in enumerate(rows[1:], start=2):
        item = {headers[i]: (row[i] if i < len(row) else None) for i in range(len(headers))}
        try:
            reminder_text = str(item.get('text') or '').strip(); when_raw = str(item.get('scheduled_local') or '').strip()
            if not reminder_text or not when_raw:
                raise ValueError('нужны text и scheduled_local')
            dt_local = parse_user_time(when_raw, user['timezone_name'])
            assigned_user_id = message.from_user.id
            assigned_raw = str(item.get('assigned_user_id') or '').strip()
            if assigned_raw:
                if not assigned_raw.isdigit() or not await get_user(int(assigned_raw)):
                    raise ValueError('assigned_user_id не найден среди активных пользователей')
                assigned_user_id = int(assigned_raw)
            await create_reminder(message.from_user.id, assigned_user_id, reminder_text, dt_local.astimezone(timezone.utc).isoformat(), str(item.get('category') or 'work').strip() or 'work', str(item.get('priority') or 'medium').strip() or 'medium', str(item.get('note') or '').strip() or None)
            imported += 1
        except Exception as e:
            errors.append(f"Строка {idx}: {e}")
    await state.clear()
    await message.answer(f"Импорт Excel завершён.\n\n✅ Импортировано: {imported}" + (f"\n⚠️ Ошибок: {len(errors)}\n" + "\n".join(errors[:10]) if errors else ""))


@router.message(CsvImportFlow.wait_file)
@router.message(ExcelImportFlow.wait_file)
async def wrong_payload(message: Message) -> None:
    await message.answer("Пришли файл документом. Чтобы выйти из режима импорта, нажми /start или выбери другой раздел.")


@router.message(F.text == "📊 Статистика")
async def stats_view(message: Message) -> None:
    data = await stats(message.from_user.id)
    await message.answer(stats_text(data), parse_mode="HTML")


@router.message(F.text == "⚙️ Настройки")
async def settings(message: Message) -> None:
    user = await get_user(message.from_user.id)
    await message.answer(f"<b>⚙️ Настройки</b>\n\n🌍 {user['timezone_name']}", parse_mode="HTML", reply_markup=settings_kb(bool(user['quiet_hours_enabled'])))


@router.message(F.text == "🆔 Мой ID")
async def myid(message: Message) -> None:
    await message.answer(f"Твой Telegram ID: <code>{message.from_user.id}</code>", parse_mode="HTML")


async def _mode_for_user(reminder: dict, user_id: int) -> str:
    if reminder.get('assigned_user_id') == user_id and reminder.get('owner_user_id') != user_id:
        return 'assigned'
    if reminder.get('owner_user_id') == user_id and reminder.get('assigned_user_id') != user_id:
        return 'owner'
    return 'shared'


async def send_window(message: Message, title: str, start_days: int, end_days: int) -> None:
    user = await get_user(message.from_user.id)
    if not user:
        return
    now_local = datetime.now(tz=to_local(datetime.now(timezone.utc), user['timezone_name']).tzinfo)
    start_local = (now_local + timedelta(days=start_days)).replace(hour=0, minute=0, second=0, microsecond=0)
    end_local = (now_local + timedelta(days=end_days)).replace(hour=0, minute=0, second=0, microsecond=0)
    rows = await list_by_time_window(message.from_user.id, start_local.astimezone(timezone.utc).isoformat(), end_local.astimezone(timezone.utc).isoformat())
    if not rows:
        await message.answer(f"<b>{title}</b>\nПока пусто.", parse_mode="HTML")
        return
    lines = [f"<b>{title}</b>"]
    for r in rows[:20]:
        dt_local = to_local(datetime.fromisoformat(r['scheduled_at_utc']), user['timezone_name'])
        owner = await get_user(r.get('owner_user_id'))
        assignee = await get_user(r.get('assigned_user_id'))
        lines.append(list_line(r, dt_local, user_label(owner, r.get('owner_user_id')), user_label(assignee, r.get('assigned_user_id')), await _mode_for_user(r, message.from_user.id)))
    await message.answer("\n\n".join(lines), parse_mode='HTML')


async def send_reminders_page(target: Message | CallbackQuery, user_id: int, page: int, category: str | None = None, priority: str | None = None) -> None:
    user = await get_user(user_id)
    reminders, total_pages = await list_active_reminders(user_id, page, category=category, priority=priority)
    subtitle = f"Фильтр: {category}" if category else (f"Фильтр: приоритет {priority}" if priority else None)
    if not reminders:
        text = "<b>📋 Активные напоминания</b>\nПока пусто."
        (await target.answer(text, parse_mode='HTML')) if isinstance(target, Message) else (await target.message.answer(text, parse_mode='HTML'), await target.answer())
        return
    header = page_header("📋 Активные напоминания", page, total_pages, subtitle)
    for idx, reminder in enumerate(reminders):
        dt_local = to_local(datetime.fromisoformat(reminder['scheduled_at_utc']), user['timezone_name'])
        owner = await get_user(reminder.get('owner_user_id'))
        assignee = await get_user(reminder.get('assigned_user_id'))
        reminder['attachments_count'] = await count_attachments(reminder['id'])
        mode = await _mode_for_user(reminder, user_id)
        text = (header + "\n\n" if idx == 0 else "") + reminder_card(reminder, dt_local, user_label(owner, reminder.get('owner_user_id')), user_label(assignee, reminder.get('assigned_user_id')), mode)
        kb = reminder_actions(reminder['id'], 'assignee' if mode == 'assigned' and not reminder.get('assignee_can_edit') else 'owner', bool(reminder.get('assignee_can_edit')), reminder.get('status'))
        if isinstance(target, Message):
            await target.answer(text, parse_mode='HTML', reply_markup=kb)
        else:
            await target.message.answer(text, parse_mode='HTML', reply_markup=kb)
    pg = pager(page, total_pages)
    if pg:
        if isinstance(target, Message):
            await target.answer("Навигация", reply_markup=pg)
        else:
            await target.message.answer("Навигация", reply_markup=pg); await target.answer()


async def send_assigned_page(target: Message | CallbackQuery, user_id: int, page: int) -> None:
    user = await get_user(user_id)
    reminders, total_pages = await list_assigned_to_me(user_id, page)
    if not reminders:
        text = "<b>📥 Мне поставили</b>\nПока пусто."
        if isinstance(target, Message): await target.answer(text, parse_mode='HTML')
        else: await target.message.answer(text, parse_mode='HTML'); await target.answer()
        return
    header = page_header("📥 Мне поставили", page, total_pages)
    for idx, reminder in enumerate(reminders):
        dt_local = to_local(datetime.fromisoformat(reminder['scheduled_at_utc']), user['timezone_name'])
        owner = await get_user(reminder.get('owner_user_id'))
        reminder['attachments_count'] = await count_attachments(reminder['id'])
        text = (header + "\n\n" if idx == 0 else "") + reminder_card(reminder, dt_local, owner_label=user_label(owner, reminder.get('owner_user_id')), mode='assigned')
        kb = reminder_actions(reminder['id'], 'assignee', bool(reminder.get('assignee_can_edit')), reminder.get('status'))
        if isinstance(target, Message): await target.answer(text, parse_mode='HTML', reply_markup=kb)
        else: await target.message.answer(text, parse_mode='HTML', reply_markup=kb)
    pg = pager(page, total_pages, prefix='assignpage')
    if pg:
        if isinstance(target, Message): await target.answer('Навигация', reply_markup=pg)
        else: await target.message.answer('Навигация', reply_markup=pg); await target.answer()


async def send_owner_page(target: Message | CallbackQuery, user_id: int, page: int) -> None:
    user = await get_user(user_id)
    reminders, total_pages = await list_created_by_me(user_id, page)
    if not reminders:
        text = "<b>📤 Я поставил</b>\nПока пусто."
        if isinstance(target, Message): await target.answer(text, parse_mode='HTML')
        else: await target.message.answer(text, parse_mode='HTML'); await target.answer()
        return
    header = page_header("📤 Я поставил", page, total_pages)
    for idx, reminder in enumerate(reminders):
        dt_local = to_local(datetime.fromisoformat(reminder['scheduled_at_utc']), user['timezone_name'])
        assignee = await get_user(reminder.get('assigned_user_id'))
        reminder['attachments_count'] = await count_attachments(reminder['id'])
        text = (header + "\n\n" if idx == 0 else "") + reminder_card(reminder, dt_local, assignee_label=user_label(assignee, reminder.get('assigned_user_id')), mode='owner')
        kb = reminder_actions(reminder['id'], 'owner', bool(reminder.get('assignee_can_edit')), reminder.get('status'))
        if isinstance(target, Message): await target.answer(text, parse_mode='HTML', reply_markup=kb)
        else: await target.message.answer(text, parse_mode='HTML', reply_markup=kb)
    pg = pager(page, total_pages, prefix='ownerpage')
    if pg:
        if isinstance(target, Message): await target.answer('Навигация', reply_markup=pg)
        else: await target.message.answer('Навигация', reply_markup=pg); await target.answer()


async def send_month_calendar(target: Message | CallbackQuery, user_id: int, year: int, month: int) -> None:
    user = await get_user(user_id)
    if not user:
        return
    tz = to_local(datetime.now(timezone.utc), user['timezone_name']).tzinfo
    start_local = datetime(year, month, 1, tzinfo=tz)
    end_local = datetime(year + (month == 12), 1 if month == 12 else month + 1, 1, tzinfo=tz)
    counts = await list_month_counts(user_id, start_local.astimezone(timezone.utc).isoformat(), end_local.astimezone(timezone.utc).isoformat())
    month_days = calendar.Calendar(firstweekday=0).monthdayscalendar(year, month)
    text = calendar_title(year, month) + "\n\n• точка рядом с днем = есть задачи"
    kb = calendar_kb(year, month, month_days, counts)
    if isinstance(target, Message): await target.answer(text, parse_mode='HTML', reply_markup=kb)
    else: await target.message.edit_text(text, parse_mode='HTML', reply_markup=kb); await target.answer()


@router.callback_query(F.data.startswith('calnav:'))
async def calendar_nav(callback: CallbackQuery) -> None:
    _, year, month, shift = callback.data.split(':')
    year = int(year); month = int(month) + int(shift)
    if month < 1: month, year = 12, year - 1
    elif month > 12: month, year = 1, year + 1
    await send_month_calendar(callback, callback.from_user.id, year, month)


@router.callback_query(F.data == 'caltoday')
async def calendar_today(callback: CallbackQuery) -> None:
    now = datetime.now(); await send_month_calendar(callback, callback.from_user.id, now.year, now.month)


@router.callback_query(F.data.startswith('calday:'))
async def calendar_day(callback: CallbackQuery) -> None:
    day_key = callback.data.split(':', 1)[1]
    user = await get_user(callback.from_user.id)
    if not user: await callback.answer(); return
    start_local = datetime.fromisoformat(day_key + 'T00:00:00').replace(tzinfo=to_local(datetime.now(timezone.utc), user['timezone_name']).tzinfo)
    end_local = start_local + timedelta(days=1)
    rows = await list_by_time_window(callback.from_user.id, start_local.astimezone(timezone.utc).isoformat(), end_local.astimezone(timezone.utc).isoformat())
    if not rows: await callback.answer('На этот день задач нет', show_alert=True); return
    lines = [f"<b>🗓 {start_local.strftime('%d.%m.%Y')}</b>"]
    for r in rows[:20]:
        dt_local = to_local(datetime.fromisoformat(r['scheduled_at_utc']), user['timezone_name'])
        owner = await get_user(r.get('owner_user_id')); assignee = await get_user(r.get('assigned_user_id'))
        r['attachments_count'] = await count_attachments(r['id'])
        lines.append(list_line(r, dt_local, user_label(owner, r.get('owner_user_id')), user_label(assignee, r.get('assigned_user_id')), await _mode_for_user(r, callback.from_user.id)))
    await callback.message.answer("\n\n".join(lines), parse_mode='HTML'); await callback.answer()


@router.callback_query(F.data.startswith('page:'))
async def page_change(callback: CallbackQuery) -> None:
    await send_reminders_page(callback, callback.from_user.id, int(callback.data.split(':', 1)[1]))


@router.callback_query(F.data.startswith('assignpage:'))
async def assign_page_change(callback: CallbackQuery) -> None:
    await send_assigned_page(callback, callback.from_user.id, int(callback.data.split(':', 1)[1]))


@router.callback_query(F.data.startswith('ownerpage:'))
async def owner_page_change(callback: CallbackQuery) -> None:
    await send_owner_page(callback, callback.from_user.id, int(callback.data.split(':', 1)[1]))


@router.callback_query(F.data == 'flt:all')
async def filter_all(callback: CallbackQuery) -> None:
    await send_reminders_page(callback, callback.from_user.id, 1)


@router.callback_query(F.data.startswith('flt:'))
async def filter_category(callback: CallbackQuery) -> None:
    await send_reminders_page(callback, callback.from_user.id, 1, category=callback.data.split(':', 1)[1])


@router.callback_query(F.data.startswith('fltprio:'))
async def filter_priority(callback: CallbackQuery) -> None:
    await send_reminders_page(callback, callback.from_user.id, 1, priority=callback.data.split(':', 1)[1])
